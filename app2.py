
import streamlit as st
import docx
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import Pt, RGBColor
from io import BytesIO
import requests
import fitz  # PyMuPDF
import pandas as pd
import re
from concurrent.futures import ThreadPoolExecutor

DEEPSEEK_API_KEY = "sk-18a6c87deb544ee6891ec8a3a5d744d6"
DEEPSEEK_API_URL = "https://api.deepseek.com/v1/chat/completions"

st.set_page_config(page_title="Startup Analyzer AI", layout="wide")
st.title("🚀 Startup Analyzer (AI-Powered Report Generator)")

uploaded_template = st.file_uploader("📄 Upload analysis template (.docx)", type=["docx"])
uploaded_files = st.file_uploader("📎 Upload supporting files (.pdf, .docx, .xlsx, .txt)", accept_multiple_files=True)

def extract_yellow_prompts(doc_file):
    doc = docx.Document(doc_file)
    prompts = []
    for para in doc.paragraphs:
        for run in para.runs:
            if run.font.highlight_color == WD_COLOR_INDEX.YELLOW:
                prompts.append(run.text.strip())
    return prompts

def extract_text_from_pdfs(pdf_paths):
    text = ""
    for file in pdf_paths:
        try:
            with fitz.open(stream=file.read(), filetype="pdf") as doc:
                for page in doc:
                    text += page.get_text()
        except:
            continue
    return text

def extract_summary_from_excel(file):
    text = ""
    try:
        xl = pd.ExcelFile(file)
        for sheet in xl.sheet_names:
            df = xl.parse(sheet)
            text += f"\n--- Sheet: {sheet} ---\n"
            text += df.head(5).to_string(index=False)
    except Exception as e:
        text += f"[Excel error: {e}]"
    return text

def build_combined_context(files):
    text = ""
    for file in files:
        name = file.name.lower()
        if name.endswith(".pdf"):
            text += extract_text_from_pdfs([file])
        elif name.endswith(".docx"):
            try:
                doc = docx.Document(file)
                for para in doc.paragraphs:
                    text += para.text + "\n"
            except:
                continue
        elif name.endswith(".xlsx"):
            text += extract_summary_from_excel(file)
        elif name.endswith(".txt"):
            text += file.read().decode(errors='ignore') + "\n"
    return text[:10000] if text else "Das Startup ist ein innovatives Projekt. Bitte analysiere es allgemein als Startup-Analyst."

def clean_markdown(text):
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    text = re.sub(r'#+\s*', '', text)
    text = re.sub(r'^\s*-\s*', '• ', text, flags=re.M)
    return text.strip()

def deepseek_generate(prompt, context):
    headers = {"Authorization": f"Bearer {DEEPSEEK_API_KEY}", "Content-Type": "application/json"}
    data = {
        "model": "deepseek-chat",
        "messages": [
            {"role": "system", "content": "Du bist ein professioneller Startup-Analyst. Schreibe formelle und prägnante Antworten auf Deutsch in maximal 1000 Zeichen."},
            {"role": "user", "content": f"{prompt}\n\nKontext:\n{context}"}
        ],
        "temperature": 0.7
    }
    try:
        res = requests.post(DEEPSEEK_API_URL, headers=headers, json=data, timeout=60)
        res.raise_for_status()
        text = res.json()["choices"][0]["message"]["content"].strip()
        fallback_phrases = ["tut mir leid", "nicht bereitgestellt", "nicht verfügbar", "leider"]
        if any(phrase in text.lower() for phrase in fallback_phrases):
            return prompt
        return clean_markdown(text[:1000])
    except:
        return prompt

def fill_doc_with_answers(template_file, answers):
    doc = docx.Document(template_file)
    idx = 0
    for para in doc.paragraphs:
        for run in para.runs:
            if run.font.highlight_color == WD_COLOR_INDEX.YELLOW and idx < len(answers):
                run.text = answers[idx]
                run.font.highlight_color = None
                run.font.name = "Calibri"
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(0, 0, 0)
                idx += 1
    return doc

if uploaded_template and st.button("🤖 Generate Report"):
    with st.spinner("🔍 Extracting prompts and analyzing context..."):
        prompts = extract_yellow_prompts(uploaded_template)
        context = build_combined_context(uploaded_files) if uploaded_files else "Das Startup ist ein innovatives Projekt. Bitte analysiere es allgemein als Startup-Analyst."

    st.info(f"🟡 Found {len(prompts)} prompts. Generating AI responses...")

    with st.spinner("🧠 Generating with DeepSeek..."):
        with ThreadPoolExecutor(max_workers=6) as executor:
            answers = list(executor.map(lambda p: deepseek_generate(p, context), prompts))

    with st.spinner("📄 Finalizing document..."):
        final_doc = fill_doc_with_answers(uploaded_template, answers)
        output = BytesIO()
        final_doc.save(output)

    st.success("✅ Your investor-ready report is ready!")
    st.download_button("📥 Download Final Report (.docx)", output.getvalue(), file_name="final_report.docx")
